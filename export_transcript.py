"""
Eksporte transkriptu daudzajos formatos:
- Klean .txt (apvienoti runataju segmenti)
- Word .docx (ar krasamem un formatesanu)
- SRT subtitri (video producentiem)
"""
from pathlib import Path
import json
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = Path(__file__).resolve().parents[1]
TEXT_DIR = BASE_DIR / "text"

SPEAKER_COLORS = [
    RGBColor(0x2E, 0x86, 0xAB),
    RGBColor(0xA0, 0x44, 0x68),
    RGBColor(0x4E, 0x9F, 0x3D),
    RGBColor(0xD0, 0x71, 0x1B),
    RGBColor(0x6A, 0x4C, 0x93),
]


def format_timestamp(seconds):
    total = int(seconds)
    h = total // 3600
    m = (total % 3600) // 60
    s = total % 60
    return f"{h:02}:{m:02}:{s:02}"


def format_srt_timestamp(seconds):
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = int(seconds % 60)
    ms = int((seconds * 1000) % 1000)
    return f"{h:02}:{m:02}:{s:02},{ms:03}"


def merge_consecutive_segments(segments):
    if not segments:
        return []

    merged = []
    current = dict(segments[0])

    for seg in segments[1:]:
        same_speaker = seg["speaker"] == current["speaker"]
        same_language = seg["language"] == current["language"]
        gap = seg["start"] - current["end"]

        if same_speaker and same_language and gap < 5:
            current["end"] = seg["end"]
            current["text"] += " " + seg["text"]
        else:
            merged.append(current)
            current = dict(seg)

    merged.append(current)
    return merged


def load_speaker_names(names_file):
    """Ielade speaker_names no faila, megina daudzus kodejumus."""
    speaker_names = {}
    content = None

    for encoding in ['utf-8-sig', 'utf-8', 'utf-16', 'cp1257']:
        try:
            content = names_file.read_text(encoding=encoding)
            break
        except UnicodeDecodeError:
            continue

    if not content:
        return speaker_names

    for line in content.strip().split("\n"):
        line = line.strip()
        if not line or line.startswith("#"):
            continue
        if "=" in line:
            key, value = line.split("=", 1)
            speaker_names[key.strip()] = value.strip()

    return speaker_names


def export_clean_txt(merged, output_path, speaker_names):
    with output_path.open("w", encoding="utf-8") as f:
        for block in merged:
            speaker = speaker_names.get(block["speaker"], block["speaker"])
            start = format_timestamp(block["start"])
            end = format_timestamp(block["end"])
            lang = block["language"].upper()

            f.write(f"[{start} - {end}] {speaker} ({lang}):\n")
            f.write(f"{block['text']}\n\n")


def export_docx(merged, output_path, source_name, speaker_names):
    doc = Document()

    title = doc.add_heading('Audio transkripts', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info_para = doc.add_paragraph()
    info_para.add_run("Avota fails: ").bold = True
    info_para.add_run(source_name)

    unique_speakers = sorted(set(b["speaker"] for b in merged))
    speakers_para = doc.add_paragraph()
    speakers_para.add_run("Runataaji: ").bold = True
    display_names = [speaker_names.get(s, s) for s in unique_speakers]
    speakers_para.add_run(f"{len(unique_speakers)} ({', '.join(display_names)})")

    doc.add_paragraph()

    speaker_color_map = {}
    for i, speaker in enumerate(unique_speakers):
        speaker_color_map[speaker] = SPEAKER_COLORS[i % len(SPEAKER_COLORS)]

    doc.add_heading('Transkripts', level=2)

    for block in merged:
        speaker = block["speaker"]
        speaker_display = speaker_names.get(speaker, speaker)
        start = format_timestamp(block["start"])
        end = format_timestamp(block["end"])
        lang = block["language"].upper()

        para = doc.add_paragraph()

        time_run = para.add_run(f"[{start} - {end}] ")
        time_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        time_run.font.size = Pt(9)

        speaker_run = para.add_run(f"{speaker_display}")
        speaker_run.bold = True
        speaker_run.font.color.rgb = speaker_color_map[speaker]

        lang_run = para.add_run(f" ({lang}):")
        lang_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        lang_run.font.size = Pt(9)

        text_para = doc.add_paragraph()
        text_para.paragraph_format.left_indent = Inches(0.3)
        text_para.add_run(block["text"])

    doc.save(output_path)


def export_srt(merged, output_path, speaker_names):
    with output_path.open("w", encoding="utf-8") as f:
        for i, block in enumerate(merged, 1):
            speaker = speaker_names.get(block["speaker"], block["speaker"])
            start = format_srt_timestamp(block["start"])
            end = format_srt_timestamp(block["end"])

            f.write(f"{i}\n")
            f.write(f"{start} --> {end}\n")
            f.write(f"{speaker}: {block['text']}\n\n")


def main():
    json_files = list(TEXT_DIR.glob("*_speakers.json"))

    if not json_files:
        print(f"Nav atrasti _speakers.json faili mape {TEXT_DIR}")
        return

    for json_file in json_files:
        print(f"\nApstrada: {json_file.name}")

        with json_file.open("r", encoding="utf-8") as f:
            data = json.load(f)

        segments = data.get("segments", [])
        source_name = data.get("source", "unknown")

        if not segments:
            print(f"  Nav segmentu, izlaiz")
            continue

        merged = merge_consecutive_segments(segments)
        print(f"  Saliek: {len(segments)} -> {len(merged)} bloku")

        base_name = json_file.stem.replace("_speakers", "")
        names_file = BASE_DIR / "audio" / f"{base_name}_names.txt"
        speaker_names = {}

        if names_file.exists():
            print(f"  Atrastas konfiguracijas vardus: {names_file.name}")
            speaker_names = load_speaker_names(names_file)
            for k, v in speaker_names.items():
                print(f"    {k} -> {v}")

        txt_path = TEXT_DIR / f"{base_name}_clean.txt"
        docx_path = TEXT_DIR / f"{base_name}_clean.docx"
        srt_path = TEXT_DIR / f"{base_name}_clean.srt"

        export_clean_txt(merged, txt_path, speaker_names)
        export_docx(merged, docx_path, source_name, speaker_names)
        export_srt(merged, srt_path, speaker_names)

        print(f"  Saglabati:")
        print(f"    - {txt_path.name}")
        print(f"    - {docx_path.name}")
        print(f"    - {srt_path.name}")

    print("\nViss pabeigts!")


if __name__ == "__main__":
    main()